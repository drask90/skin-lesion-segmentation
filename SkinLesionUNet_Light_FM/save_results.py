# -*- coding: utf-8 -*-
"""
Evaluate all 5 fold models on their respective validation splits,
then save quantitative + qualitative results to a Word document.
"""

import os
os.environ["KMP_DUPLICATE_LIB_OK"] = "TRUE"

import numpy as np
import torch
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from data_preparation import get_fold_loaders
from model import UNet


# ─────────────────────────────────────────────────────────────
# Evaluation
# ─────────────────────────────────────────────────────────────
def evaluate_fold(fold, device):
    weight_path = f"best_model_fold{fold + 1}.pth"
    if not os.path.exists(weight_path):
        print(f"  Weight not found: {weight_path}")
        return None

    model = UNet().to(device)
    model.load_state_dict(torch.load(weight_path, map_location=device))
    model.eval()

    _, val_loader = get_fold_loaders(fold)

    iou_total = dsc_total = sen_total = spe_total = acc_total = 0.0

    with torch.no_grad():
        for images, masks in val_loader:
            masks = masks.to(device)
            outputs = torch.sigmoid(model(images.to(device)))
            preds = (outputs > 0.5).float()

            tp = (preds * masks).sum(dim=(1, 2, 3))
            tn = ((1 - preds) * (1 - masks)).sum(dim=(1, 2, 3))
            fp = (preds * (1 - masks)).sum(dim=(1, 2, 3))
            fn = ((1 - preds) * masks).sum(dim=(1, 2, 3))

            iou_total += (tp / (tp + fp + fn + 1e-8)).mean().item()
            dsc_total += (2 * tp / (2 * tp + fp + fn + 1e-8)).mean().item()
            sen_total += (tp / (tp + fn + 1e-8)).mean().item()
            spe_total += (tn / (tn + fp + 1e-8)).mean().item()
            acc_total += ((tp + tn) / (tp + tn + fp + fn + 1e-8)).mean().item()

    n = len(val_loader)
    return {
        "DSC":         dsc_total / n,
        "IoU":         iou_total / n,
        "Sensitivity": sen_total / n,
        "Specificity": spe_total / n,
        "Accuracy":    acc_total / n,
    }


# ─────────────────────────────────────────────────────────────
# Word document builder
# ─────────────────────────────────────────────────────────────
def build_word_doc(fold_results, qualitative_image_path, out_path):
    doc = Document()

    # ── Title ─────────────────────────────────────────────────
    title = doc.add_heading("Skin Lesion Segmentation – Results Report", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(
        "Dataset: ISIC 2016  |  Model: UNet (base=64)  |  "
        "5-Fold Cross-Validation"
    ).alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    # ── Training Setup ─────────────────────────────────────────
    doc.add_heading("1. Loss Function & Hyperparameters", level=1)

    doc.add_heading("Loss Function", level=2)
    doc.add_paragraph(
        "A combined BCE + IoU loss was used during training:"
    )
    doc.add_paragraph(
        "Loss = BCE(output, target) + (1 − IoU(output, target))",
        style="List Bullet",
    )
    doc.add_paragraph(
        "where IoU = (output · target) / (output + target − output · target + ε),  ε = 1e-8"
    )

    doc.add_heading("Hyperparameters", level=2)
    hp_table = doc.add_table(rows=1, cols=2)
    hp_table.style = "Table Grid"
    hdr = hp_table.rows[0].cells
    for i, h in enumerate(["Hyperparameter", "Value"]):
        hdr[i].text = h
        hdr[i].paragraphs[0].runs[0].bold = True
        hdr[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    hyperparams = [
        ("Model",            "UNet (base=16)"),
        ("Input size",       "256 × 256"),
        ("Optimizer",        "AdamW"),
        ("Learning rate",    "1e-4"),
        ("Epochs",           "200"),
        ("Batch size",       "8"),
        ("Cross-validation", "5-Fold (KFold, shuffle=True, random_state=42)"),
        ("Augmentation",     "Random horizontal flip, random rotation ±15°"),
        ("Normalisation",    "mean=[0.485, 0.456, 0.406], std=[0.229, 0.224, 0.225]"),
    ]
    for param, value in hyperparams:
        row = hp_table.add_row().cells
        row[0].text = param
        row[1].text = value

    doc.add_paragraph()

    # ── Quantitative Results ───────────────────────────────────
    doc.add_heading("2. Quantitative Results", level=1)

    metrics = ["DSC", "IoU", "Sensitivity", "Specificity", "Accuracy"]
    headers = ["Fold"] + metrics

    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"

    # Header row
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        hdr[i].paragraphs[0].runs[0].bold = True
        hdr[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Data rows
    valid_results = []
    for fold, res in enumerate(fold_results):
        row = table.add_row().cells
        row[0].text = f"Fold {fold + 1}"
        row[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        if res is not None:
            valid_results.append(res)
            for i, m in enumerate(metrics):
                row[i + 1].text = f"{res[m]:.4f}"
                row[i + 1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        else:
            for i in range(len(metrics)):
                row[i + 1].text = "N/A"
                row[i + 1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Mean ± Std row
    if valid_results:
        row = table.add_row().cells
        row[0].text = "Mean ± Std"
        row[0].paragraphs[0].runs[0].bold = True
        row[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        for i, m in enumerate(metrics):
            vals = [r[m] for r in valid_results]
            row[i + 1].text = f"{np.mean(vals):.4f} ± {np.std(vals):.4f}"
            row[i + 1].paragraphs[0].runs[0].bold = True
            row[i + 1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    # Mean summary paragraph
    if valid_results:
        doc.add_heading("Summary", level=2)
        for m in metrics:
            vals = [r[m] for r in valid_results]
            doc.add_paragraph(
                f"{m}: {np.mean(vals):.4f} ± {np.std(vals):.4f}",
                style="List Bullet",
            )

    doc.add_paragraph()

    # ── Qualitative Results ────────────────────────────────────
    doc.add_heading("3. Qualitative Results", level=1)
    doc.add_paragraph(
        "The figure below shows 5 representative samples. Each row displays "
        "the input image, ground-truth mask, and the model prediction overlaid "
        "with colour-coded error regions:"
    )
    legend = doc.add_paragraph()
    legend.add_run("Yellow").font.color.rgb = RGBColor(0xCC, 0xCC, 0x00)
    legend.add_run(" = True Positive (TP)   ")
    legend.add_run("Red").font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
    legend.add_run(" = False Negative (FN – missed)   ")
    legend.add_run("Green").font.color.rgb = RGBColor(0x00, 0xAA, 0x00)
    legend.add_run(" = False Positive (FP – over-segmented)")

    if os.path.exists(qualitative_image_path):
        doc.add_picture(qualitative_image_path, width=Inches(6.0))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        doc.add_paragraph(
            f"[Image not found: {qualitative_image_path}]"
        )

    doc.save(out_path)
    print(f"Saved: '{out_path}'")


# ─────────────────────────────────────────────────────────────
# Main
# ─────────────────────────────────────────────────────────────
def main():
    device = torch.device("cuda" if torch.cuda.is_available() else "cpu")
    print(f"Device: {device}\n")

    fold_results = []
    for fold in range(5):
        print(f"Evaluating Fold {fold + 1}/5 ...")
        res = evaluate_fold(fold, device)
        fold_results.append(res)
        if res:
            print(f"  DSC={res['DSC']:.4f}  IoU={res['IoU']:.4f}  "
                  f"Sen={res['Sensitivity']:.4f}  Spe={res['Specificity']:.4f}  "
                  f"Acc={res['Accuracy']:.4f}")

    build_word_doc(
        fold_results,
        qualitative_image_path="qualitative_comparison_isic2016.png",
        out_path="results_report.docx",
    )


if __name__ == "__main__":
    main()
